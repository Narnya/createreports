#! /usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font, Run
import fnmatch
import yaml
import csv
import codecs
import os
import re
import natsort
import matplotlib.pyplot as plt
import numpy
import numpy.core._dtype_ctypes
#print(os.path.join(os.path.dirname(os.path.abspath(__file__)),'dictreading.yaml'))

#def getConfig():
    #with open(os.path.join(os.path.dirname(os.path.abspath(__file__)),'dictreading.yaml'), 'r', encoding='utf-8') as fileObject:
        #try:
            #return yaml.load(fileObject)
        #except yaml.YAMLError as exception:
            #print(exception)
        # config = data
        #config['keywords'] = data['keywords']
        #config[keywordscsv = data['keywordscsv']
        #config[path1] = data['path1']
        #config[path2] = data['path2']
        #config[path3] = data['path3']
       
config = {
    "availablereportnames": [
        "Report1",
        "Report2",
        "Report3",
        "Report5"
    ],
    "path1": "C:\\Users\\M.Gusev\\Michail\\Skoltech\\2018\\2018_08_12_Tupolev\\2019_01_14_CFX_Script\\example_matrix_of_experiments\\Report1\\",
    "path2": "C:\\Users\\M.Gusev\\Michail\\Skoltech\\2018\\2018_08_12_Tupolev\\2018_12_21_Python_Script\\pro_v3\\",
    "path3": "C:\\Users\\M.Gusev\\Michail\\Skoltech\\2018\\2018_08_12_Tupolev\\2019_01_14_CFX_Script\\example_matrix_of_experiments\\Report1\\",
    "tableoutput": "No",
    "keywords": {
        "Report1": {
            "chapter1": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от угла атаки",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter2": {
                "title": "Зависимость коэффициента лобового сопротивления от угла атаки",
                "description": "Зависимость коэффициента лобового сопротивления от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter3": {
                "title": "Зависимость коэффициента аэродинамической поперечной силы от угла атаки",
                "description": "Зависимость коэффициента аэродинамической поперечной силы от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter4": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от угла атаки",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter5": {
                "title": "Зависимость коэффициента аэродинамического момента крена от угла атаки",
                "description": "Зависимость коэффициента аэродинамического момента крена от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter6": {
                "title": "Зависимость коэффициента аэродинамического момента рыскания от угла атаки при beta = 2.11 и mach = 1.32",
                "description": "Зависимость коэффициента аэродинамического момента рыскания от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter7": {
                "title": "Поляра 1-го рода",
                "description": "Поляра 1-го рода при beta = {beta} и мах = {mach}"
            },
            "chapter8": {
                "title": "Зависимость аэродинамического качества от угла атаки",
                "description": "Зависимость аэродинамического качества от угла атаки при beta = {beta} и мах = {mach}"
            },
            "chapter9": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от коэффициента аэродинамической подъемной силы",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от коэффициента аэродинамической подъемной силы при beta = {beta} и мах = {mach}"
            },
            "chapter10": {
                "title": "Зависимость аэродинамического качества от коэффициента аэродинамической подъемной силы",
                "description": "Зависимость аэродинамического качества от коэффициента аэродинамической подъемной силы при beta = {beta} и мах = {mach}"
            },
            "chapter11": {
                "title": "Зависимость коэффициента лобового сопротивления при нулевой подъемной силе от числа Маха",
                "description": "Зависимость коэффициента лобового сопротивления при нулевой подъемной силе от числа Маха при beta = {beta}"
            },
            "chapter12": {
                "title": "Зависимость минимального значения коэффициента лобового сопротивления от числа Маха",
                "description": "Зависимость минимального значения коэффициента лобового сопротивления от числа Маха при beta = {beta}"
            },
            "chapter13": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от коэффициента индуктивного сопротивления",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от коэффициента индуктивного сопротивления при beta = {beta} и мах = {mach}"
            },
            "chapter14": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от числа Маха",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter15": {
                "title": "Зависимость коэффициента лобового сопротивления от числа Маха",
                "description": "Зависимость коэффициента лобового сопротивления от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter16": {
                "title": "Зависимость коэффициента аэродинамической поперечной силы от числа Маха",
                "description": "Зависимость коэффициента аэродинамической поперечной силы от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter17": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от числа Маха",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter18": {
                "title": "Зависимость коэффициента аэродинамического момента крена от числа Маха",
                "description": "Зависимость коэффициента аэродинамического момента крена от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter19": {
                "title": "Зависимость коэффициента аэродинамического момента рыскания от числа Маха",
                "description": "Зависимость коэффициента аэродинамического момента рыскания от числа Маха при alpha = {alpha} и beta = {beta}"
            },
            "chapter20": {
                "title": "Зависимость координаты положения фокуса самолета от числа Маха",
                "description": "Зависимость координаты положения фокуса самолета от числа Маха при beta = {beta}"
            },
            "chapter21": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от угла скольжения",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от угла скольжения при alpha = {alpha} и beta = {beta}"
            },
            "chapter22": {
                "title": "Зависимость коэффициента лобового сопротивления от числа Маха",
                "description": "Зависимость коэффициента лобового сопротивления от числа Маха при alpha = {alpha} и мах = {mach}"
            },
            "chapter23": {
                "title": "Зависимость коэффициента аэродинамической поперечной силы от угла скольжения",
                "description": "Зависимость коэффициента аэродинамической поперечной силы от угла скольжения при alpha = {alpha} и мах = {mach}"
            },
            "chapter24": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от угла скольжения",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от угла скольжения при alpha = {alpha} и мах = {mach}"
            },
            "chapter25": {
                "title": "Зависимость коэффициента аэродинамического момента крена от угла скольжения",
                "description": "Зависимость коэффициента аэродинамического момента крена от угла скольжения при alpha = {alpha} и мах = {mach}"
            },
            "chapter26": {
                "title": "Зависимость коэффициента аэродинамического момента рыскания от угла скольжения",
                "description": "Зависимость коэффициента аэродинамического момента рыскания от угла скольжения при alpha = {alpha} и мах = {mach}"
            },
            "chapter27": {
                "title": "Распределение коэффициента давления по плоскостям",
                "description": "Распределение коэффициента давления по плоскости, перпендикулярной оси {axis} при {axis}={value}, alpha = {alpha}, beta = {beta} и мах = {mach}"
            }
        },
        "Report2": {
            "chapter1": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от времени",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от времени"
            },
            "chapter2": {
                "title": "Зависимость коэффициента лобового сопротивления от времени",
                "description": "Зависимость коэффициента лобового сопротивления от времени"
            },
            "chapter3": {
                "title": "Зависимость коэффициента аэродинамической поперечной силы от времени",
                "description": "Зависимость коэффициента аэродинамической поперечной силы от времени"
            },
            "chapter4": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от времени",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от времени"
            },
            "chapter5": {
                "title": "Зависимость коэффициента аэродинамического момента крена от времени",
                "description": "Зависимость коэффициента аэродинамического момента крена от времени"
            },
            "chapter6": {
                "title": "Зависимость коэффициента аэродинамического момента рыскания от времени",
                "description": "Зависимость коэффициента аэродинамического момента рыскания от времени"
            },
            "chapter7": {
                "title": "Зависимость угла атаки от времени",
                "description": "Зависимость угла атаки от времени"
            },
            "chapter8": {
                "title": "Зависимость угла скольжения от времени",
                "description": "Зависимость угла скольжения от времени"
            },
            "chapter9": {
                "title": "Зависимость угла крена от времени",
                "description": "Зависимость угла крена от времени"
            },
            "chapter10": {
                "title": "Зависимость координаты центра масс по оси ОХ от времени",
                "description": "Зависимость координаты центра масс по оси ОХ от времени"
            },
            "chapter11": {
                "title": "Зависимость координаты центра масс по оси ОY от времени",
                "description": "Зависимость координаты центра масс по оси ОY от времени"
            },
            "chapter12": {
                "title": "Зависимость координаты центра масс по оси ОZ от времени",
                "description": "Зависимость координаты центра масс по оси ОZ от времени"
            }
        },
        "Report3": {
            "chapter1": {
                "title": "Зависимость приращения коэффициента аэродинамической подъемной силы от угла атаки",
                "description": "Зависимость приращения коэффициента аэродинамической подъемной силы от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter2": {
                "title": "Зависимость приращения коэффициента лобового сопротивления от угла атаки",
                "description": "Зависимость приращения коэффициента лобового сопротивления от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter3": {
                "title": "Зависимость приращения коэффициента аэродинамической поперечной силы от угла атаки",
                "description": "Зависимость приращения коэффициента аэродинамической поперечной силы от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter4": {
                "title": "Зависимость приращения коэффициента аэродинамического момента тангажа от угла атаки",
                "description": "Зависимость приращения коэффициента аэродинамического момента тангажа от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter5": {
                "title": "Зависимость приращения коэффициента аэродинамического момента крена от угла атаки",
                "description": "Зависимость приращения коэффициента аэродинамического момента крена от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter6": {
                "title": "Зависимость приращения коэффициента аэродинамического момента рыскания от угла атаки",
                "description": "Зависимость приращения коэффициента аэродинамического момента рыскания от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter7": {
                "title": "Зависимость приращения аэродинамического качества от угла атаки",
                "description": "Зависимость приращения аэродинамического качества от угла атаки при beta = {beta} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter8": {
                "title": "Зависимость приращения коэффициента аэродинамической подъемной силы от угла скольжения",
                "description": "Зависимость приращения коэффициента аэродинамической подъемной силы от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter9": {
                "title": "Зависимость приращения коэффициента лобового сопротивления от угла скольжения",
                "description": "Зависимость приращения коэффициента лобового сопротивления от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter10": {
                "title": "Зависимость приращения коэффициента аэродинамической поперечной силы от угла скольжения",
                "description": "Зависимость приращения коэффициента аэродинамической поперечной силы от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter11": {
                "title": "Зависимость приращения коэффициента аэродинамического момента тангажа от угла скольжения",
                "description": "Зависимость приращения коэффициента аэродинамического момента тангажа от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter12": {
                "title": "Зависимость приращения коэффициента аэродинамического момента крена от угла скольжения",
                "description": "Зависимость приращения коэффициента аэродинамического момента крена от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter13": {
                "title": "Зависимость приращения коэффициента аэродинамического момента рыскания от угла скольжения",
                "description": "Зависимость приращения коэффициента аэродинамического момента рыскания от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter14": {
                "title": "Зависимость приращения аэродинамического качества от угла скольжения",
                "description": "Зависимость приращения аэродинамического качества от угла скольжения при alpha = {alpha} и мах = {mach} для {axis}-ой конфигурации"
            },
            "chapter15": {
                "title": "Приращение коэффициента аэродинамической подъемной силы в зависимости от положения надстройки",
                "description": "Приращение коэффициента аэродинамической подъемной силы в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter16": {
                "title": "Приращение коэффициента лобового сопротивления в зависимости от положения надстройки",
                "description": "Приращение коэффициента лобового сопротивления в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter17": {
                "title": "Приращение коэффициента аэродинамической поперечной силы в зависимости от положения надстройки",
                "description": "Приращение коэффициента аэродинамической поперечной силы в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter18": {
                "title": "Приращения коэффициента аэродинамического момента тангажа в зависимости от положения надстройки",
                "description": "Приращения коэффициента аэродинамического момента тангажа в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter19": {
                "title": "Приращение коэффициента аэродинамического момента крена от в зависимости от положения надстройки",
                "description": "Приращение коэффициента аэродинамического момента крена от в зависимости от положения надстройки alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter20": {
                "title": "Приращение коэффициента аэродинамического момента рыскания в зависимости от положения надстройки",
                "description": "Приращение коэффициента аэродинамического момента рыскания в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter21": {
                "title": "Приращение аэродинамического качества в зависимости от положения надстройки",
                "description": "Приращение аэродинамического качества в зависимости от положения надстройки при alpha = {alpha}, beta = {beta} и мах = {mach}"
            },
            "chapter22": {
                "title": "Распределение коэффициента давления по плоскостям",
                "description": "Распределение коэффициента давления по плоскости, перпендикулярной оси {axis} при {axis}={value}, alpha = {alpha}, beta = {beta} и мах = {mach} для {n}-ой конфигурации"
            }
        },
        "Report5": {
            "chapter1": {
                "title": "Зависимость коэффициента аэродинамической подъемной силы от угла атаки",
                "description": "Зависимость коэффициента аэродинамической подъемной силы от угла атаки при мах = {mach}"
            },
            "chapter2": {
                "title": "Зависимость коэффициента аэродинамической поперечной силы от угла атаки",
                "description": "Зависимость коэффициента аэродинамической поперечной силы от угла атаки при мах = {mach}"
            },
            "chapter3": {
                "title": "Зависимость коэффициента аэродинамического момента тангажа от угла атаки",
                "description": "Зависимость коэффициента аэродинамического момента тангажа от угла атаки при мах = {mach}"
            },
            "chapter4": {
                "title": "Зависимость коэффициента аэродинамического момента крена от угла атаки",
                "description": "Зависимость коэффициента аэродинамического момента крена от угла атаки при мах = {mach}"
            },
            "chapter5": {
                "title": "Зависимость коэффициента аэродинамического момента рыскания от угла атаки",
                "description": "Зависимость коэффициента аэродинамического момента рыскания от угла атаки при мах = {mach}"
            },
            "chapter6": {
                "title": "Комплекс производных коэффициента нормальной аэродинамической силы",
                "description": "Комплекс производных коэффициента нормальной аэродинамической силы при мах = {mach}"
            },
            "chapter7": {
                "title": "Комплекс динамических производных коэффициента момента тангажа в фазе с угловой скоростью по оси OZ",
                "description": "Комплекс динамических производных коэффициента момента тангажа в фазе с угловой скоростью по оси OZ при мах = {mach}"
            },
            "chapter8": {
                "title": "Комплекс динамических производных коэффициента боковой аэродинамической силы в фазе с угловой скоростью по оси OY",
                "description": "Комплекс динамических производных коэффициента боковой аэродинамической силы в фазе с угловой скоростью по оси OY при мах = {mach}"
            },
            "chapter9": {
                "title": "Комплекс динамических производных коэффициента момента крена в фазе с угловой скоростью по оси OY",
                "description": "Комплекс динамических производных коэффициента момента крена в фазе с угловой скоростью по оси OY при мах = {mach}"
            },
            "chapter10": {
                "title": "Комплекс динамических производных коэффициента момента рыскания в фазе с угловой скоростью по оси OY",
                "description": "Комплекс динамических производных коэффициента момента рыскания в фазе с угловой скоростью по оси OY при мах = {mach}"
            }
        }
    }
}
# пути для отчета
pathRep0 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Report')
pathRep1 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Report1')
pathRep2 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Report2')
pathRep3 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Report3')
pathRep5 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Report5')

print([pathRep0, pathRep1, pathRep2, pathRep3, pathRep5])
# нахождение доступных путей
if os.path.exists(pathRep1):
    pathRep = pathRep1 + "\\"
elif os.path.exists(pathRep2):
    pathRep = pathRep2 + "\\"
elif os.path.exists(pathRep3):
    pathRep = pathRep3 + "\\"
elif os.path.exists(pathRep5):
    pathRep = pathRep5 + "\\"
elif os.path.exists(pathRep0):
    pathRep = pathRep0 + "\\"

config['path1'] = pathRep
config['path2'] = os.path.dirname(os.path.abspath(__file__)) + "\\"
config['path3'] = pathRep

# keywords = data("{items:{keywords:}")
print('Keywords: "' + '", "'.join(config['availablereportnames']) + '".')

# получение частей имени таблицы
def getTableFileParts(tableFileName):
    return tableFileName.split('.csv')[0].split('_')

# парсинг всех названий рисунков с расширением .png и возврат в отсортированном виде
def findImageFileNames():
    imageFileNames = []
    for imageFileName in os.listdir(config['path1']):
        if fnmatch.fnmatch(imageFileName, '*.png'):
            imageFileNames.append(imageFileName)
    return natsort.natsorted(imageFileNames)

print(findImageFileNames())

imagename = findImageFileNames()
# print(imagename[0])
# получение названий таблиц по частям
def getTableParts(tableFileName):
    return natsort.natsorted(tableFileName.split('.csv')[0].split('_'))

# получение названий рисунков по частям
def getImageParts(imageFileName):
    return imageFileName.split('.png')[0].split('_')

# нахождение названий таблиц с расширением .csv
def findTableFileNames():
    tableFileNames = []
    for tableFileName in os.listdir(config['path3']):
        if fnmatch.fnmatch(tableFileName, '*.csv'):
            tableFileNames.append(tableFileName)
    return tableFileNames

# t=findTableFileNames()
# print(t)

# чтение таблицы
def importTable(tableFileName):
    with open(config['path3'] + tableFileName, 'r', newline="\n", encoding='utf-8') as fileObject:
        reader = csv.reader(x.replace('\0', ',') for x in fileObject)
        return list(reader)

# добавление таблицы в документ
def addTable(tableData, document):
    table = document.add_table(rows=len(tableData), cols=len(tableData[0])) # add style
    table.style = 'Table Grid'
    for rowIndex, row in enumerate(tableData):
        for columnIndex, col in enumerate(row):
            cell = table.cell(rowIndex, columnIndex)
            cell.text = col
            cell.width = Inches(3.5)
            #print(cell)

# центрирование
def center(document):
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# выравнивание по правому краю
def right(document):
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# выравнивание по левому краю
def left(document):
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# создание изображения
def create_images():
    for table_file_name in findTableFileNames():
        x = []
        y = []
        # чтение таблицы
        with open(pathRep + table_file_name, 'r') as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            headers = next(csv_reader)
            if len(headers) != 2:
                continue
            isHeadersOmitted = False
            # проверка наличия названия осей
            if headers[0].isdigit() and headers[1].isdigit(): # если шапки - цифры, значит названия осей опущены
                isHeadersOmitted = True # флаг опущения названий осей, т.е. шапки таблиц
                x.append(float(headers[0]))
                y.append(float(headers[1]))
            # заполнение массива значений x и y
            for row in csv_reader:
                x.append(float(row[0]))
                y.append(float(row[1]))
            # если названия опущены, тогда название осей задается по умолчанию
            if isHeadersOmitted:
                plt.plot(x, y, marker='o')
                plt.xlabel('X')
                plt.ylabel('Y')
            # иначе, заполняется первой строчкой таблицы, шапка таблицы - это название осей
            else:
                # print(x)
                # print(y)
                # plt.plot(x, y, label=table_file_name.replace('.csv', ''))
                # plt.style.use('ggplot')
                plt.plot(x, y, label='{0}({1})'.format(headers[1], headers[0]), marker='o')
                # plt.axvline(x=0, color='k')
                # plt.axhline(y=0, color='k')
                plt.xlabel(headers[0] + (', °' if headers[0] in ('alpha', 'betta') else ''))
                plt.ylabel(headers[1])
                # plt.title('XY plot')
            plt.legend()
            plt.grid(True)
        # plt.show()
        plt.savefig(pathRep + table_file_name.replace('.csv', '.png'))
        plt.clf()

# создание отчета
def createReports():
    create_images() # генерирование графиков
    tableFileNamePart0 = '' # значение первой части названия таблицы - "Report№", отвечает за вид отчета
    tableFileNamePart1 = '' # значение второй части названия таблицы - "Cxy", отвечает за изменение главы внутри одного отчета
    imagePart2 = ''
    imagePart3 = ''
    paragraphIndex = 0 # номер параграфа внутри отчета
    imageIndex = 0 # номер рисунка внутри главы
    tableIndex = 0 # номер таблицы внутри главы
    isFirstImage = None
    document = None
    # главный цикл - вход в перебор картинок из сгенерированных
    for imageFileName in findImageFileNames():
        imageParts = getImageParts(imageFileName) # массив частей названия изображения
        print(imageParts, imageParts[0])
        # print(imageFileName)
        if imageParts[0] not in config['availablereportnames']:  # проверка на наличие части названия рисунка с доступным именем в словаре
            continue
        if tableFileNamePart0 != imageParts[0]: # если название таблицы не совпадает с названием изображения, отслеживается изменение названия отчета
            paragraphIndex = 0 # определяем начальный индекс параграфа
            tableIndex = 0 # определяем начальный индекс таблицы
            print('Here', tableFileNamePart0)
            if document is not None:
                document.save(config['path2'] + tableFileNamePart0 + '.docx') # создание нового документа если выявили изменение в первой части названия рисунка
            tableFileNamePart0 = imageParts[0] # установка нового названия отчета
            document = Document()
            imageIndex = 0 # обнуление индекса картинки для нового отчета
        if tableFileNamePart1 != imageParts[1]: # условие изменения главы внутри одного отчета
            tableFileNamePart1 = imageParts[1] # меняем номер главы
            #imagePart3 = imageParts[3]
            #print(tableFileNamePart1)
            paragraphIndex += 1 # изменение индекса параграфа
            paragraph = document.add_paragraph() # добавление нового параграфа
            chapterNumber = int(re.findall('^\d+', imageParts[1])[0])  # извлекаем номер главы рисунка
            paragraph.add_run('Глава '+ str(paragraphIndex) + '.' + ' ' + config['keywords'][imageParts[0]]['chapter' + str(chapterNumber)]['title']).bold = True #add part3 добавление названия главы, нумерация глав
            print(paragraphIndex)
            center(document) # центрирование
            isFirstImage = True # флаг для первой картинки отчета
        imageFilePath = config['path1'] + str(imageFileName[:-3]) + 'png' # формирование названия рисунка для добавления в документ
        if os.path.isfile(imageFilePath): # если такой файл существует
            imageIndex += 1
            document.add_picture(imageFilePath, width=Inches(5.0), height=Inches(3.5)) # добавление картинки
            center(document)
            # print(len(imageParts))
            chapterNumber = int(re.findall('^\d+', imageParts[1])[0]) # извлекаем номер главы рисунка
            print('chapterNumber', chapterNumber)
            if paragraphIndex != 0 and imageParts[0] in config['availablereportnames']: # проверка наличия имение отчета в словаре
                prefix = 'Рисунок ' + str(imageIndex) + '. ' # нумерация рисунка
                description = config['keywords'][imageParts[0]]['chapter' + str(chapterNumber)]['description'] # подпись рисунка: описание
                postfix = '.'
            parameters = None
            if imageParts[0] == 'Report1': # заполнение первого отчета
                # условие вхождения номера главы в заданный диапазон
                if chapterNumber in list([1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 13]):
                    parameters = {'beta': imageParts[3], 'mach': imageParts[4]}

                if chapterNumber in list([14, 15, 16, 17, 18, 19, 21]):
                    parameters = {'alpha': imageParts[3], 'beta': imageParts[4]}

                if chapterNumber in list([11, 12, 20]):
                    parameters = {'beta': imageParts[3]}

                if chapterNumber in list([22, 23, 24, 25, 26]):
                    parameters = {'alpha': imageParts[3], 'mach': imageParts[4]}

                if imageParts[1] == '27Cp':
                    print('true')
                    parameters = {'axis': imageParts[2], 'value': imageParts[3], 'alpha': imageParts[4],
                                  'beta': imageParts[5], 'mach': imageParts[6]}
                print(imageParts[1]) # 13Cxi(Cya)
                print(paragraphIndex) # выдает сквозную нумерацию 11
                print(imageParts[3]) # 0.0
                print('Hello!', prefix, description, postfix)
                document.add_paragraph(prefix + description.format(**parameters) + postfix) # добавление названия графика для Report1
                center(document)
                document.add_paragraph()
                center(document)

            if imageParts[0] == 'Report2':
                document.add_paragraph(prefix + description) # добавление названия графика для Report2
                document.add_paragraph()
                center(document)

            if imageParts[0] == 'Report3':
                #
                print(chapterNumber)
                if chapterNumber in range(1, 7 + 1):
                    #print(chapterNumber)
                    parameters = {'axis': imageParts[2], 'beta': imageParts[3], 'mach': imageParts[4]}
                if chapterNumber in range(8, 14 + 1):
                    parameters = {'axis': imageParts[2], 'alpha': imageParts[3], 'mach': imageParts[4]}
                if chapterNumber in range(15, 21 + 1):
                    parameters = {'alpha': imageParts[3], 'beta': imageParts[4], 'mach': imageParts[5]}
                if chapterNumber == 22:
                    parameters = {'n': imageParts[3], 'axis': imageParts[2], 'value': imageParts[4], 'alpha': imageParts[5],
                                  'beta': imageParts[6], 'mach': imageParts[7]}
                document.add_paragraph(prefix + description.format(**parameters) + postfix) # добавление названия графика для Report3
                document.add_paragraph()
                print(description)
                center(document)

            if imageParts[0] == 'Report5':
                print(imageParts[3])
                print(prefix)
                print(description)
                document.add_paragraph(prefix + description.format(mach=imageParts[3])) # добавление названия графика для Report5
                document.add_paragraph()
                center(document)
        else:
            document.add_paragraph("Данных для построения этого отчёта было недостаточно.")
            document.add_page_break()
            continue

        # если это первое изображение
        if isFirstImage:
            if (str(imageFileName[:-3]) + 'csv') in findTableFileNames(): # проверка существования таблицы
                tableData = importTable(str(imageFileName[:-3]) + 'csv') # чтение таблицы
                tableIndex += 1
                print('Tabl' ,imageFileName, tableIndex, tableData) #  Это печатается 2 раза, как???
                if chapterNumber != 0 and imageParts[0] in config['availablereportnames']: # если существует параграф
                    # формирование названия таблицы
                    prefixtabl = 'Таблица ' + str(tableIndex) + '. '
                    description = config['keywords'][imageParts[0]]['chapter' + str(chapterNumber)][
                        'description']
                    postfix = '.'
                    print('prefixtable', prefixtabl)
                if imageParts[0] == 'Report1':
                    if chapterNumber in list([1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 13]):
                        document.add_paragraph(
                            prefixtabl + description.format(beta=imageParts[3], mach=imageParts[4]) + postfix) # добавление названия таблицы
                        print('Yes1')

                    if chapterNumber in list([14, 15, 16, 17, 18, 19, 21]):
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4]) + postfix)
                        print('Yes2')

                    if chapterNumber in list([11, 12, 20]):
                        document.add_paragraph(
                            prefixtabl + description.format(beta=imageParts[3]) + postfix)
                        print('Yes3')

                    if chapterNumber in list([22, 23, 24, 25, 26]):
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], mach=imageParts[4]) + postfix)
                        print('Yes4')

                    if imageParts[1] == '27Cp':
                        # добавление названия таблицы
                        if imageParts[2] == 'x':
                            document.add_paragraph(
                                prefixtabl + description.format(axis=imageParts[2], value=imageParts[3],
                                                                alpha=imageParts[4],
                                                                beta=imageParts[5], mach=imageParts[6]) + postfix)
                        if imageParts[2] == 'y':
                            document.add_paragraph(
                                prefixtabl + description.format(axis=imageParts[2], value=imageParts[3],
                                                                alpha=imageParts[4],
                                                                beta=imageParts[5], mach=imageParts[6]) + postfix)
                        if imageParts[2] == 'z':
                            document.add_paragraph(
                                prefixtabl + description.format(axis=imageParts[2], value=imageParts[3],
                                                                alpha=imageParts[4],
                                                                beta=imageParts[5], mach=imageParts[6]) + postfix)

                    addTable(tableData, document)
                    left(document)
                    document.add_page_break()
            # возможность отключения вывода таблицы
            if config['tableoutput'] == 'Yes':
                if imageParts[0] == 'Report2':
                    document.add_paragraph(prefixtabl + description)
                    center(document)
                if imageParts[0] == 'Report3':
                    if imageParts[1] == '1dCya(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '2dCxa(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '3dCz(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '4dmza(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '5dmxa(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '6dmya(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '7dK(alpha)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], alpha=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '8dCya(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '9dCxa(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '10dCz(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '11dmza(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                        #print(imageParts[0])
                    if imageParts[1] == '12dmxa(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '13dmya(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '14dK(betta)':
                        document.add_paragraph(
                            prefixtabl + description.format(axis=imageParts[2], beta=imageParts[3],
                                                            mach=imageParts[4]) + postfix)
                    if imageParts[1] == '15dCya(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '16dCxa(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '17dCz(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '18dmza(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '19dmxa(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                        print(imageParts[0])
                    if imageParts[1] == '20dmya(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '21dK(x,y,z)':
                        document.add_paragraph(
                            prefixtabl + description.format(alpha=imageParts[3], beta=imageParts[4],
                                                            mach=imageParts[5]) + postfix)
                    if imageParts[1] == '22Cp':
                        if imageParts[2] == 'x':
                            document.add_paragraph(
                                prefixtabl + description.format(n=imageParts[3], axis=imageParts[2], value=imageParts[4],
                                                                alpha=imageParts[5],
                                                                beta=imageParts[6],
                                                                mach=imageParts[7]) + postfix)
                        if imageParts[2] == 'y':
                            document.add_paragraph(
                                prefixtabl + description.format(axis=imageParts[2], value=imageParts[3],
                                                                alpha=imageParts[4],
                                                                beta=imageParts[5],
                                                                mach=imageParts[6]) + postfix)
                        if imageParts[2] == 'z':
                            document.add_paragraph(
                                prefixtabl + description.format(axis=imageParts[2], value=imageParts[3],
                                                                alpha=imageParts[4],
                                                                beta=imageParts[5],
                                                                mach=imageParts[6]) + postfix)
                    center(document)

                if imageParts[0] == 'Report5':
                    document.add_paragraph(prefixtabl + description.format(mach=imageParts[3]))
                    center(document)

                center(document)
                addTable(tableData, document)
                document.add_page_break()

            # else:
            #     if imageParts[0] == 'Report1':
            #         center(document)
            #         addTable(tableData, document)
            #         document.add_page_break()
            #         print('tabledataa', tableData)
        isFirstImage = False
    if document:
        document.save(config['path2'] + tableFileNamePart0 + '.docx')
createReports()

